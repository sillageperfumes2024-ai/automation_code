import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
from datetime import datetime
import argparse
import win32com.client
from pypdf import PdfWriter
import shutil
import logging
import time


def create_shipping_label(order_data, order_number):
    """
    Create a Word document with shipping label for a single order
    """
    # Create a new document
    doc = Document()
    
    # Set up the document margins (smaller margins for label)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.2)
        section.bottom_margin = Inches(0.2)
        section.left_margin = Inches(0.2)
        section.right_margin = Inches(0.2)
    
    # Add title
    title = doc.add_heading('Sillage Perfumes', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add a line break
    doc.add_paragraph()
    
    # Create a table for the shipping label layout
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = None#'Table Grid'
    
    # Set column widths
    table.columns[0].width = Inches(7)
    #table.columns[1].width = Inches(7)
    
    # Get the cells
    receiver_cell = table.cell(0, 0)
    sender_cell = table.cell(1, 0)
    
    
    # SENDER INFORMATION
    sender_paragraph = sender_cell.paragraphs[0]
    sender_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add sender header
    sender_run = sender_paragraph.add_run("FROM:\n")
    sender_run.bold = True
    sender_run.font.size = Pt(18)
    
    # Add sender details
    sender_details = """Sillage Perfumes
Outer Ring Road Doddanekundi
Bengaluru, 560037
Karnataka, India
Phone: 8904620890"""
    
    sender_info_run = sender_paragraph.add_run(sender_details)
    sender_info_run.font.size = Pt(16)
    
    # RECEIVER INFORMATION
    receiver_paragraph = receiver_cell.paragraphs[0]
    receiver_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add receiver header
    receiver_run = receiver_paragraph.add_run("TO:\n")
    receiver_run.bold = True
    receiver_run.font.size = Pt(24)
    
    # Get receiver details from order data
    
    shipping_name = order_data['Shipping Name'].iloc[0] if not pd.isna(order_data['Shipping Name'].iloc[0]) else "N/A"
    shipping_street = order_data['Shipping Street'].iloc[0] if not pd.isna(order_data['Shipping Street'].iloc[0]) else "N/A"
    shipping_city = order_data['Shipping City'].iloc[0] if not pd.isna(order_data['Shipping Street'].iloc[0]) else "N/A"
    shipping_zip = str(order_data['Shipping Zip'].iloc[0] if not pd.isna(order_data['Shipping Street'].iloc[0]) else "N/A").replace(".0","")
    shipping_province = order_data['Shipping Province'].iloc[0] if not pd.isna(order_data['Shipping Street'].iloc[0]) else "N/A"
    shipping_country = order_data['Shipping Country'].iloc[0] if not pd.isna(order_data['Shipping Street'].iloc[0]) else "N/A"
    shipping_phone = str(order_data['Shipping Phone'].iloc[0] if not pd.isna(order_data['Shipping Street'].iloc[0]) else "N/A").replace(".0","")
    
    receiver_details = f"{shipping_name}\n{shipping_street}\n{shipping_city}, {shipping_zip}\n{shipping_province},{shipping_country}\nPhone:{shipping_phone}"
    receiver_info_run = receiver_paragraph.add_run(receiver_details)
    receiver_info_run.font.size = Pt(22)
    
    # Add space between sections
    #doc.add_paragraph()
    
    # Order Information Section
    #order_info = doc.add_heading('ORDER DETAILS', level=2)
    #order_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Order details table
    order_table = doc.add_table(rows=1, cols=1)
    order_table.style = None #'Table Grid'
    order_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Order Number
    order_table.cell(0, 0).text = f"Order Number: {str(order_number)}"
    #order_table.cell(0, 1).text = str(order_number)
    
    # Date
    #order_table.cell(1, 0).text = "Date:"
    #order_table.cell(1, 1).text = datetime.now().strftime("%Y-%m-%d")
    
    # Make headers bold
    for i in range(1):
        order_table.cell(i, 0).paragraphs[0].runs[0].bold = True
    
    # Add space
    #doc.add_paragraph()
    # Group items by name and count quantities
    item_counts = {}
    for _, row in order_data.iterrows():
        item_name = row['Lineitem name']
        if pd.notna(item_name):
            if item_name in item_counts:
                item_counts[item_name] += 1
            else:
                item_counts[item_name] = 1

    # Calculate total number of items
    total_items = sum(item_counts.values())
    
    # Add Total Items table
    total_items_table = doc.add_table(rows=1, cols=2)
    total_items_table.style = None
    total_items_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Add total items information
    total_cells = total_items_table.rows[0].cells
    total_cells[0].text = 'Total Items'
    total_cells[1].text = str(total_items)
    
    # Create items table
    items_table = doc.add_table(rows=1, cols=2)
    items_table.style = None#'Table Grid'
    items_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Add header row
    hdr_cells = items_table.rows[0].cells
    hdr_cells[0].text = 'Item Name'
    hdr_cells[1].text = 'Quantity'
    
    # Make header bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    
    # Add items to table
    max_items = 5
    item_counter = 0
    for item_name, quantity in item_counts.items():
        item_counter += 1
        if item_counter <= max_items:
            row_cells = items_table.add_row().cells
            row_cells[0].text = str(item_name)
            row_cells[1].text = str(quantity)
        else:
            break
    
    return doc

def process_order_export(file_path, output_folder, need_single_pdf=True):
    """
    Process the order export file and create shipping labels for each order
    """
    try:
        # Read the CSV file
        df = pd.read_csv(file_path)
        
        # Check if required columns exist
        required_columns = ['Name', 'Shipping Name', 'Shipping Street', 'Lineitem name', 'Shipping City', 'Shipping Zip', 'Shipping Province', 'Shipping Country', 'Shipping Phone']
        missing_columns = [col for col in required_columns if col not in df.columns]
        
        if missing_columns:
            print(f"Error: Missing columns: {missing_columns}")
            print(f"Available columns: {list(df.columns)}")
            return
        
        # Create output directory
        
        output_dir = output_folder
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # Group by order number (Name column)
        grouped_orders = df.groupby('Name')
        
        print(f"Processing {len(grouped_orders)} orders...")
        
        # Process each order
        for order_number, order_data in grouped_orders:
            try:
                # Create shipping label document
                doc = create_shipping_label(order_data, order_number)
                
                # Generate filename (sanitize order number for filename)
                safe_order_number = str(order_number).replace('/', '_').replace('\\', '_')
                filename = f"shipping_label_{safe_order_number}.docx"
                filepath = os.path.join(output_dir, filename)
                
                # Save the document
                doc.save(filepath)
                print(f"Created shipping label: {filename}")
                
            except Exception as e:
                print(f"Error processing order {order_number}: {str(e)}")
        
        print(f"\nAll shipping labels created successfully in '{output_dir}' directory!")
        
    except FileNotFoundError:
        print(f"Error: File '{file_path}' not found.")
    except Exception as e:
        print(f"Error processing file: {str(e)}")

def combine_docs_with_word(doc_folder, output_pdf):
    """
    Finds all .doc files in a folder, converts them to PDF using MS Word,
    and merges them into a single PDF file. Each original document will
    naturally start on a new page in the final PDF.
    """
    # Create a temporary folder to store the intermediate PDFs.
    temp_folder = os.path.join(doc_folder, "temp_pdfs_for_conversion")
    if os.path.exists(temp_folder):
        shutil.rmtree(temp_folder)
    os.makedirs(temp_folder)

    word_instance = None
    try:
        # --- Part 1: Convert .doc files to .pdf ---
        logging.info("Starting conversion from .doc to .pdf using Microsoft Word...")
        
        # Start a Word application instance.
        word_instance = win32com.client.Dispatch("Word.Application")
        # Keep Word hidden in the background.
        word_instance.Visible = False

        doc_files = sorted([f for f in os.listdir(doc_folder) if f.lower().endswith(".docx")])
        
        if not doc_files:
            logging.warning(f"No .doc files were found in '{doc_folder}'.")
            return

        for filename in doc_files:
            # Construct full paths for input and output files.
            doc_path = os.path.join(doc_folder, filename)
            pdf_path = os.path.join(temp_folder, f"{os.path.splitext(filename)[0]}.pdf")
            
            logging.info(f"Converting '{filename}' to PDF...")
            
            # Open the document.
            doc = word_instance.Documents.Open(doc_path)
            # The FileFormat constant for PDF is 17.
            doc.SaveAs(pdf_path, FileFormat=17)
            doc.Close()

        logging.info("All .doc files have been converted to PDF successfully.")

        # --- Part 2: Merge the generated PDFs ---
        logging.info("Merging individual PDF files...")
        pdf_merger = PdfWriter()
        
        pdf_files = [os.path.join(temp_folder, f"{os.path.splitext(f)[0]}.pdf") for f in doc_files]

        for pdf_path in pdf_files:
            if os.path.exists(pdf_path):
                pdf_merger.append(pdf_path)
                logging.info(f"Appended '{os.path.basename(pdf_path)}' to the final document.")

        # Write the merged PDF to the output file.
        with open(output_pdf, "wb") as out_file:
            pdf_merger.write(out_file)
        
        pdf_merger.close()
        logging.info(f"Successfully created the combined PDF: {output_pdf}")

    except Exception as e:
        logging.error(f"An error occurred: {e}")
        logging.error("Please ensure Microsoft Word is installed and you have permissions to access the folders.")
        
    finally:
        # --- Part 3: Clean up ---
        # Ensure the Word application is closed, even if an error occurred.
        if word_instance:
            word_instance.Quit()
        
        # Remove the temporary folder with intermediate PDFs.
        if os.path.exists(temp_folder):
            logging.info("Cleaning up temporary files...")
            shutil.rmtree(temp_folder)
        
        logging.info("Process finished.")


if __name__ == "__main__":
    
    logging.basicConfig(level=logging.DEBUG,  # Change to INFO, WARNING, ERROR, CRITICAL as needed
    format="%(asctime)s - %(levelname)s - %(message)s")

    # Read file from argument
    parser = argparse.ArgumentParser()
    parser.add_argument('--orders_file', type=str, help='absolute path of the orders from Shopify')
    parser.add_argument('--output_folder', type=str, help='absolute path of the output folder')
    parser.add_argument('--save_single_pdf', type=bool, default=True, help='Need to save a single pdf file')

    # Parse the arguments
    args = parser.parse_args()
    # Specify your input file path
    
    input_file = args.orders_file #"order_export.csv"  # Change this to your actual file path
    
    print("Shipping Label Generator")
    print("=" * 30)
    
    # Check if file exists
    if not os.path.exists(input_file):
        print(f"Please make sure '{input_file}' exists in the current directory.")
        print("Or update the 'input_file' variable with the correct path.")
    else:
        current_date = datetime.now().strftime("%d_%m_%Y_%H_%M")
        
        output_dir = os.path.abspath(os.path.join(args.output_folder, current_date))
        output_pdf_file = os.path.abspath(os.path.join(output_dir, "combined_orders.pdf"))
        
        process_order_export(input_file, output_dir)
        time.sleep(5)
        print(f"Output folder: {output_dir}")
        print(f"Output pdf file: {output_pdf_file}")
        if args.save_single_pdf:
            combine_docs_with_word(output_dir, output_pdf_file)